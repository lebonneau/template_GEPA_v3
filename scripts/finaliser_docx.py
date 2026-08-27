#!/usr/bin/env python3
"""Ajoute la page titre GÉPA, une table des matières et la pagination au DOCX Pandoc."""

from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt


def read_macro(text: str, name: str, default: str = "") -> str:
    match = re.search(rf"\\{re.escape(name)}\{{([^}}]*)\}}", text)
    return match.group(1).strip() if match else default


def add_field(paragraph, instruction: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate, text, end):
        run._r.append(element)


def move_before(paragraph, anchor) -> None:
    anchor.addprevious(paragraph._p)


def set_child_value(parent, tag: str, attr: str, value: str):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    child.set(qn(attr), value)
    return child


def normalize_tables(doc: Document) -> None:
    usable_twips = 9405
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        cols = max(1, len(table.columns))
        if cols == 3:
            widths = [4405, 2500, 2500]
        else:
            base = usable_twips // cols
            widths = [base] * cols
            widths[-1] += usable_twips - sum(widths)

        tbl_pr = table._tbl.tblPr
        tbl_w = set_child_value(tbl_pr, "w:tblW", "w:w", str(usable_twips))
        tbl_w.set(qn("w:type"), "dxa")
        set_child_value(tbl_pr, "w:tblInd", "w:w", "0").set(qn("w:type"), "dxa")
        set_child_value(tbl_pr, "w:jc", "w:val", "center")

        grid = table._tbl.tblGrid
        for child in list(grid):
            grid.remove(child)
        for width in widths:
            col = OxmlElement("w:gridCol")
            col.set(qn("w:w"), str(width))
            grid.append(col)

        for row_index, row in enumerate(table.rows):
            for col_index, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                tc_pr = cell._tc.get_or_add_tcPr()
                tc_w = set_child_value(tc_pr, "w:tcW", "w:w", str(widths[col_index]))
                tc_w.set(qn("w:type"), "dxa")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT if col_index == 0 else WD_ALIGN_PARAGRAPH.CENTER
                    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    paragraph.paragraph_format.space_after = Pt(0)
                    for run in paragraph.runs:
                        run.bold = row_index == 0


def main() -> None:
    if len(sys.argv) != 4:
        raise SystemExit("Usage: finaliser_docx.py corps.docx informations.tex sortie.docx")

    body_path, info_path, output_path = map(Path, sys.argv[1:])
    info = info_path.read_text(encoding="utf-8")
    values = {
        "titre": read_macro(info, "GEPATitre", "Titre du travail"),
        "soustitre": read_macro(info, "GEPASousTitre"),
        "auteur": read_macro(info, "GEPAAuteur", "Prénom Nom"),
        "destinataire": read_macro(info, "GEPADestinataire"),
        "cours": read_macro(info, "GEPACours"),
        "lieu": read_macro(info, "GEPALieu"),
        "date": read_macro(info, "GEPADate"),
    }

    doc = Document(body_path)
    normalize_tables(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_field(footer_p, " PAGE ", "2")

    first_body_paragraph = doc.paragraphs[0] if doc.paragraphs else None
    if first_body_paragraph is not None:
        first_body_paragraph.paragraph_format.page_break_before = True
    first = first_body_paragraph._p if first_body_paragraph is not None else doc._body._element.sectPr

    for paragraph in doc.paragraphs:
        if paragraph.text.strip().startswith("Source :"):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            paragraph.paragraph_format.space_before = Pt(4)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                run.font.size = Pt(10)
    created = []

    def new_para(text: str = "", style: str | None = None):
        p = doc.add_paragraph()
        if style:
            style_obj = next((item for item in doc.styles if item.name == style), None)
            if style_obj is None:
                raise KeyError(f"Style DOCX introuvable : {style}")
            p.style = style_obj
        if text:
            p.add_run(text)
        created.append(p)
        return p

    new_para("École de politique appliquée", "GEPA Institution")
    new_para("Faculté des lettres et sciences humaines", "GEPA Institution")
    new_para("Université de Sherbrooke", "GEPA Institution")
    new_para(values["titre"], "GEPA Titre")
    if values["soustitre"]:
        new_para(values["soustitre"], "GEPA Sous-titre")
    new_para("Par " + values["auteur"], "GEPA Auteur")
    new_para(values["destinataire"], "GEPA Destinataire")
    new_para(values["cours"], "GEPA Cours")
    new_para(values["lieu"], "GEPA Lieu")
    p_date = new_para(values["date"], "GEPA Date")
    p_date.add_run().add_break()
    p_date.runs[-1]._r[-1].set(qn("w:type"), "page")

    toc_heading = new_para("Table des matières", "Heading 1")
    toc_para = new_para(style="Normal")
    add_field(toc_para, ' TOC \\o "1-3" \\h \\z \\u ', "Mettez à jour ce champ dans Word (clic droit > Mettre à jour le champ).")
    for p in created:
        move_before(p, first)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)


if __name__ == "__main__":
    main()
